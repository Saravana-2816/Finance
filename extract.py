# requirements.txt

# schema.py
"""
Pydantic schemas for transaction validation and output formatting.
"""
from typing import List, Optional
from pydantic import BaseModel, Field, validator
from datetime import datetime
import re

class Transaction(BaseModel):
    """Individual transaction record with validation."""
    from_field: str = Field(..., alias="from", description="Source of the transaction")
    to: str = Field(..., description="Destination/recipient of the transaction")
    credit_account: str = Field(..., description="Account being credited")
    debit_account: str = Field(..., description="Account being debited")
    expense_type: str = Field(..., description="Primary expense category")
    subcategory: str = Field(..., description="Expense subcategory")
    debit_amount: float = Field(0.0, ge=0, description="Amount debited")
    credit_amount: float = Field(0.0, ge=0, description="Amount credited")
    transaction_date: str = Field(..., description="Transaction date in YYYY-MM-DD format")
    
    @validator('transaction_date')
    def validate_date_format(cls, v):
        """Ensure date is in YYYY-MM-DD format."""
        try:
            datetime.strptime(v, '%Y-%m-%d')
            return v
        except ValueError:
            raise ValueError('Date must be in YYYY-MM-DD format')
    
    @validator('debit_amount', 'credit_amount', pre=True)
    def clean_amounts(cls, v):
        """Clean and validate monetary amounts."""
        if isinstance(v, str):
            # Remove currency symbols, commas, and whitespace
            cleaned = re.sub(r'[^\d.-]', '', v.strip())
            if not cleaned or cleaned == '-':
                return 0.0
            try:
                return float(cleaned)
            except ValueError:
                return 0.0
        return float(v) if v is not None else 0.0
    
    class Config:
        populate_by_name = True
        json_encoders = {
            float: lambda v: round(v, 2)
        }

class TransactionList(BaseModel):
    """List of validated transactions."""
    transactions: List[Transaction]
    total_count: int = Field(..., description="Total number of transactions")
    processing_metadata: Optional[dict] = Field(None, description="Processing information")
    
    @validator('total_count')
    def validate_count(cls, v, values):
        """Ensure count matches actual transaction list length."""
        if 'transactions' in values:
            actual_count = len(values['transactions'])
            if v != actual_count:
                return actual_count
        return v

# file_loader.py
"""
File loading utilities for different document formats.
"""
import os
import logging
from typing import Union, Tuple, List
import pandas as pd
import pdfplumber
from docx import Document
from pathlib import Path

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileLoader:
    """Handles loading of various file formats."""
    
    SUPPORTED_EXTENSIONS = {'.csv', '.xlsx', '.xls', '.pdf', '.docx'}
    
    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)
    
    def load_file(self, file_path: str) -> Tuple[str, Union[pd.DataFrame, str, List[str]]]:
        """
        Load file based on extension and return (file_type, content).
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (file_type, content) where content varies by type
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        self.logger.info(f"Loading {file_ext} file: {file_path}")
        
        try:
            if file_ext == '.csv':
                return self._load_csv(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._load_excel(file_path)
            elif file_ext == '.pdf':
                return self._load_pdf(file_path)
            elif file_ext == '.docx':
                return self._load_docx(file_path)
        except Exception as e:
            self.logger.error(f"Error loading file {file_path}: {str(e)}")
            raise
    
    def _load_csv(self, file_path: str) -> Tuple[str, pd.DataFrame]:
        """Load CSV file with robust encoding detection."""
        encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']
        
        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, encoding=encoding, low_memory=False)
                self.logger.info(f"Successfully loaded CSV with {encoding} encoding")
                return 'csv', df
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode CSV file with any of the tried encodings: {encodings}")
    
    def _load_excel(self, file_path: str) -> Tuple[str, pd.DataFrame]:
        """Load Excel file, trying all sheets."""
        try:
            # Try to read all sheets and combine
            excel_file = pd.ExcelFile(file_path)
            
            # If only one sheet, return it directly
            if len(excel_file.sheet_names) == 1:
                df = pd.read_excel(file_path, sheet_name=0)
                return 'excel', df
            
            # Multiple sheets - look for the one most likely to contain transactions
            transaction_keywords = ['transaction', 'statement', 'activity', 'history']
            
            for sheet_name in excel_file.sheet_names:
                if any(keyword in sheet_name.lower() for keyword in transaction_keywords):
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    self.logger.info(f"Using sheet: {sheet_name}")
                    return 'excel', df
            
            # If no obvious sheet found, use the first one
            df = pd.read_excel(file_path, sheet_name=0)
            self.logger.info(f"Using first sheet: {excel_file.sheet_names[0]}")
            return 'excel', df
            
        except Exception as e:
            raise ValueError(f"Error reading Excel file: {str(e)}")
    
    def _load_pdf(self, file_path: str) -> Tuple[str, List[str]]:
        """Extract text from PDF using pdfplumber."""
        pages_text = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        pages_text.append(text)
                        self.logger.info(f"Extracted text from page {i+1}")
            
            if not pages_text:
                self.logger.warning("No text extracted from PDF - may need OCR")
            
            return 'pdf', pages_text
            
        except Exception as e:
            raise ValueError(f"Error reading PDF file: {str(e)}")
    
    def _load_docx(self, file_path: str) -> Tuple[str, str]:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))
            
            combined_text = '\n'.join(full_text)
            self.logger.info(f"Extracted {len(full_text)} text blocks from DOCX")
            
            return 'docx', combined_text
            
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {str(e)}")

# preprocess.py
"""
Preprocessing utilities for cleaning and normalizing extracted data.
"""
import re
import pandas as pd
from typing import List, Dict, Any, Union
from datetime import datetime
import logging
from dateutil import parser

logger = logging.getLogger(__name__)

class DataPreprocessor:
    """Handles preprocessing of extracted data from various sources."""
    
    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)
        self.date_patterns = [
            r'\d{4}-\d{2}-\d{2}',           # YYYY-MM-DD
            r'\d{2}/\d{2}/\d{4}',           # MM/DD/YYYY
            r'\d{2}-\d{2}-\d{4}',           # MM-DD-YYYY
            r'\d{2}\.\d{2}\.\d{4}',         # MM.DD.YYYY
            r'\d{1,2}/\d{1,2}/\d{4}',       # M/D/YYYY
            r'\d{1,2}-\d{1,2}-\d{4}',       # M-D-YYYY
        ]
        
        self.amount_pattern = r'[-+]?\$?\d{1,3}(?:,\d{3})*(?:\.\d{2})?'
        
    def preprocess_structured_data(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Preprocess structured data (CSV/Excel).
        
        Args:
            df: Raw DataFrame
            
        Returns:
            Cleaned DataFrame
        """
        self.logger.info(f"Preprocessing structured data with {len(df)} rows")
        
        # Remove completely empty rows
        df = df.dropna(how='all')
        
        # Convert column names to lowercase and strip whitespace
        df.columns = df.columns.astype(str).str.lower().str.strip()
        
        # Remove rows that appear to be headers (repeated column names)
        header_mask = df.apply(
            lambda row: any(str(val).lower().strip() in df.columns for val in row if pd.notna(val)), 
            axis=1
        )
        df = df[~header_mask]
        
        # Identify potential transaction rows
        df = self._identify_transaction_rows(df)
        
        # Clean and standardize data types
        df = self._clean_structured_data(df)
        
        self.logger.info(f"After preprocessing: {len(df)} rows remain")
        return df
    
    def preprocess_text_data(self, text_data: Union[str, List[str]]) -> List[str]:
        """
        Preprocess unstructured text data (PDF/DOCX).
        
        Args:
            text_data: Raw text or list of text blocks
            
        Returns:
            List of cleaned text lines
        """
        if isinstance(text_data, str):
            text_data = [text_data]
        
        all_lines = []
        for text_block in text_data:
            lines = text_block.split('\n')
            all_lines.extend(lines)
        
        # Clean and filter lines
        cleaned_lines = []
        for line in all_lines:
            line = line.strip()
            if self._is_potential_transaction_line(line):
                cleaned_lines.append(line)
        
        self.logger.info(f"Preprocessed text data: {len(cleaned_lines)} potential transaction lines")
        return cleaned_lines
    
    def _identify_transaction_rows(self, df: pd.DataFrame) -> pd.DataFrame:
        """Identify rows that likely contain transaction data."""
        # Look for rows with date and amount patterns
        transaction_mask = pd.Series([False] * len(df))
        
        for col in df.columns:
            # Check for date patterns
            date_mask = df[col].astype(str).str.contains('|'.join(self.date_patterns), na=False)
            
            # Check for amount patterns
            amount_mask = df[col].astype(str).str.contains(self.amount_pattern, na=False)
            
            # Row is likely a transaction if it has both date and amount
            transaction_mask |= (date_mask & df.apply(
                lambda row: any(re.search(self.amount_pattern, str(val)) for val in row), axis=1
            ))
        
        return df[transaction_mask]
    
    def _clean_structured_data(self, df: pd.DataFrame) -> pd.DataFrame:
        """Clean structured data types and formats."""
        # Remove currency symbols and commas from potential amount columns
        for col in df.columns:
            if df[col].dtype == 'object':
                # Check if column contains amounts
                sample_values = df[col].dropna().astype(str).head(10)
                if any(re.search(self.amount_pattern, val) for val in sample_values):
                    df[col] = df[col].astype(str).str.replace(r'[\$,]', '', regex=True)
                    df[col] = pd.to_numeric(df[col], errors='ignore')
        
        return df
    
    def _is_potential_transaction_line(self, line: str) -> bool:
        """Determine if a text line potentially contains transaction data."""
        if not line or len(line.strip()) < 10:
            return False
        
        # Skip obvious header/footer patterns
        skip_patterns = [
            r'^page \d+',
            r'^statement period',
            r'^account summary',
            r'^balance forward',
            r'^opening balance',
            r'^closing balance',
            r'^total',
            r'^\*+',
            r'^-+$',
            r'^=+$',
        ]
        
        if any(re.search(pattern, line.lower()) for pattern in skip_patterns):
            return False
        
        # Must contain at least a date pattern and an amount pattern
        has_date = any(re.search(pattern, line) for pattern in self.date_patterns)
        has_amount = re.search(self.amount_pattern, line)
        
        return has_date and has_amount
    
    def normalize_date(self, date_str: str) -> str:
        """
        Normalize date string to YYYY-MM-DD format.
        
        Args:
            date_str: Date string in various formats
            
        Returns:
            Normalized date string in YYYY-MM-DD format
        """
        if not date_str or pd.isna(date_str):
            return ""
        
        try:
            # Try parsing with dateutil parser first (most flexible)
            parsed_date = parser.parse(str(date_str), fuzzy=True)
            return parsed_date.strftime('%Y-%m-%d')
        except:
            # If that fails, try manual pattern matching
            for pattern in self.date_patterns:
                match = re.search(pattern, str(date_str))
                if match:
                    date_part = match.group()
                    try:
                        if '-' in date_part:
                            parts = date_part.split('-')
                        elif '/' in date_part:
                            parts = date_part.split('/')
                        elif '.' in date_part:
                            parts = date_part.split('.')
                        else:
                            continue
                        
                        if len(parts) == 3:
                            # Determine format based on first part length
                            if len(parts[0]) == 4:  # YYYY-MM-DD
                                return f"{parts[0]}-{parts[1].zfill(2)}-{parts[2].zfill(2)}"
                            else:  # MM-DD-YYYY or DD-MM-YYYY
                                year = parts[2] if len(parts[2]) == 4 else f"20{parts[2]}"
                                return f"{year}-{parts[0].zfill(2)}-{parts[1].zfill(2)}"
                    except:
                        continue
        
        self.logger.warning(f"Could not normalize date: {date_str}")
        return ""
    
    def clean_amount(self, amount_str: Union[str, float, int]) -> float:
        """
        Clean and normalize monetary amounts.
        
        Args:
            amount_str: Amount in various formats
            
        Returns:
            Cleaned float amount
        """
        if pd.isna(amount_str) or amount_str == "":
            return 0.0
        
        if isinstance(amount_str, (int, float)):
            return float(amount_str)
        
        # Remove currency symbols, commas, and whitespace
        cleaned = re.sub(r'[^\d.-]', '', str(amount_str).strip())
        
        if not cleaned or cleaned == '-':
            return 0.0
        
        try:
            return float(cleaned)
        except ValueError:
            self.logger.warning(f"Could not parse amount: {amount_str}")
            return 0.0

# extractor.py
"""
Main extraction logic for converting preprocessed data into transaction records.
"""
import re
import pandas as pd
from typing import List, Dict, Any, Union, Optional
import logging
from preprocess import DataPreprocessor

logger = logging.getLogger(__name__)

class TransactionExtractor:
    """Extracts transaction records from preprocessed data."""
    
    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)
        self.preprocessor = DataPreprocessor()
        
        # Common column mappings for structured data
        self.column_mappings = {
            'date': ['date', 'transaction_date', 'trans_date', 'posting_date', 'effective_date'],
            'description': ['description', 'desc', 'memo', 'details', 'transaction_details', 'payee'],
            'amount': ['amount', 'transaction_amount', 'trans_amount'],
            'debit': ['debit', 'debit_amount', 'withdrawal', 'outgoing', 'payment'],
            'credit': ['credit', 'credit_amount', 'deposit', 'incoming', 'receipt'],
            'balance': ['balance', 'running_balance', 'account_balance'],
            'reference': ['reference', 'ref', 'transaction_id', 'check_number'],
        }
    
    def extract_from_structured_data(self, df: pd.DataFrame) -> List[Dict[str, Any]]:
        """
        Extract transactions from structured data (CSV/Excel).
        
        Args:
            df: Preprocessed DataFrame
            
        Returns:
            List of transaction dictionaries
        """
        self.logger.info(f"Extracting transactions from structured data ({len(df)} rows)")
        
        # Map columns to standard names
        column_map = self._map_columns(df.columns.tolist())
        
        transactions = []
        for idx, row in df.iterrows():
            try:
                transaction = self._extract_transaction_from_row(row, column_map)
                if transaction:
                    transactions.append(transaction)
            except Exception as e:
                self.logger.warning(f"Error processing row {idx}: {str(e)}")
                continue
        
        self.logger.info(f"Extracted {len(transactions)} transactions from structured data")
        return transactions
    
    def extract_from_text_data(self, lines: List[str]) -> List[Dict[str, Any]]:
        """
        Extract transactions from text data (PDF/DOCX).
        
        Args:
            lines: List of preprocessed text lines
            
        Returns:
            List of transaction dictionaries
        """
        self.logger.info(f"Extracting transactions from text data ({len(lines)} lines)")
        
        transactions = []
        for line_num, line in enumerate(lines):
            try:
                transaction = self._extract_transaction_from_line(line)
                if transaction:
                    transactions.append(transaction)
            except Exception as e:
                self.logger.warning(f"Error processing line {line_num}: {str(e)}")
                continue
        
        self.logger.info(f"Extracted {len(transactions)} transactions from text data")
        return transactions
    
    def _map_columns(self, columns: List[str]) -> Dict[str, str]:
        """Map DataFrame columns to standard field names."""
        column_map = {}
        columns_lower = [col.lower().strip() for col in columns]
        
        for standard_field, possible_names in self.column_mappings.items():
            for col_idx, col_name in enumerate(columns_lower):
                if any(possible_name in col_name for possible_name in possible_names):
                    column_map[standard_field] = columns[col_idx]  # Use original case
                    break
        
        self.logger.info(f"Column mapping: {column_map}")
        return column_map
    
    def _extract_transaction_from_row(self, row: pd.Series, column_map: Dict[str, str]) -> Optional[Dict[str, Any]]:
        """Extract transaction data from a DataFrame row."""
        # Get mapped values
        date_val = row.get(column_map.get('date', ''), '')
        description_val = row.get(column_map.get('description', ''), '')
        amount_val = row.get(column_map.get('amount', ''), 0)
        debit_val = row.get(column_map.get('debit', ''), 0)
        credit_val = row.get(column_map.get('credit', ''), 0)
        reference_val = row.get(column_map.get('reference', ''), '')
        
        # Clean and normalize values
        transaction_date = self.preprocessor.normalize_date(str(date_val))
        if not transaction_date:
            return None
        
        # Determine amounts
        if amount_val and not debit_val and not credit_val:
            # Single amount column - determine if debit or credit based on sign
            clean_amount = self.preprocessor.clean_amount(amount_val)
            if clean_amount < 0:
                debit_amount = abs(clean_amount)
                credit_amount = 0.0
            else:
                debit_amount = 0.0
                credit_amount = clean_amount
        else:
            debit_amount = self.preprocessor.clean_amount(debit_val)
            credit_amount = self.preprocessor.clean_amount(credit_val)
        
        # Skip if no meaningful amount
        if debit_amount == 0 and credit_amount == 0:
            return None
        
        # Build transaction record
        description_clean = str(description_val).strip() if description_val else "Unknown Transaction"
        
        transaction = {
            'raw_description': description_clean,
            'reference': str(reference_val).strip() if reference_val else '',
            'debit_amount': debit_amount,
            'credit_amount': credit_amount,
            'transaction_date': transaction_date,
            'source_data': 'structured'
        }
        
        return transaction
    
    def _extract_transaction_from_line(self, line: str) -> Optional[Dict[str, Any]]:
        """Extract transaction data from a text line."""
        # Pattern for common bank statement line formats
        # Typical format: DATE DESCRIPTION AMOUNT
        patterns = [
            # Pattern 1: Date at start, amount at end
            r'(\d{1,2}[-/\.]\d{1,2}[-/\.]\d{2,4})\s+(.+?)\s+([-+]?\$?\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s*$',
            # Pattern 2: Date, description, debit, credit
            r'(\d{1,2}[-/\.]\d{1,2}[-/\.]\d{2,4})\s+(.+?)\s+([-+]?\$?\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s+([-+]?\$?\d{1,3}(?:,\d{3})*(?:\.\d{2})?)',
            # Pattern 3: More flexible pattern
            r'(\d{4}[-/\.]\d{1,2}[-/\.]\d{1,2})\s+(.+?)\s+([-+]?\$?\d{1,3}(?:,\d{3})*(?:\.\d{2})?)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, line)
            if match:
                groups = match.groups()
                
                if len(groups) >= 3:
                    date_str = groups[0]
                    description = groups[1].strip()
                    amount_str = groups[2]
                    
                    # Normalize date
                    transaction_date = self.preprocessor.normalize_date(date_str)
                    if not transaction_date:
                        continue
                    
                    # Clean amount
                    amount = self.preprocessor.clean_amount(amount_str)
                    if amount == 0:
                        continue
                    
                    # Determine if debit or credit
                    if amount < 0 or '-' in amount_str:
                        debit_amount = abs(amount)
                        credit_amount = 0.0
                    else:
                        debit_amount = 0.0
                        credit_amount = amount
                    
                    # Handle case with separate debit/credit columns
                    if len(groups) >= 4:
                        credit_str = groups[3]
                        credit_amount = self.preprocessor.clean_amount(credit_str)
                        debit_amount = amount  # First amount is debit
                    
                    return {
                        'raw_description': description,
                        'reference': '',
                        'debit_amount': debit_amount,
                        'credit_amount': credit_amount,
                        'transaction_date': transaction_date,
                        'source_data': 'text'
                    }
        
        return None

# categorizer.py
"""
Transaction categorization using NLP and rule-based approaches.
"""
import re
import logging
from typing import Dict, List, Tuple, Set
import spacy
from rapidfuzz import fuzz

logger = logging.getLogger(__name__)

class TransactionCategorizer:
    """Categorizes transactions into expense types and subcategories."""
    
    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)
        
        # Load spaCy model for NER
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            self.logger.warning("spaCy model 'en_core_web_sm' not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
        
        # Vendor mapping dictionary
        self.vendor_mappings = {
            # Food & Dining
            'MCD': 'McDonald\'s',
            'MCDONALDS': 'McDonald\'s',
            'KFC': 'KFC',
            'STARBUCKS': 'Starbucks',
            'SUBWAY': 'Subway',
            'PIZZA HUT': 'Pizza Hut',
            'DOMINOS': 'Domino\'s Pizza',
            'BURGER KING': 'Burger King',
            'TACO BELL': 'Taco Bell',
            'CHIPOTLE': 'Chipotle',
            
            # Retail
            'WALMART': 'Walmart',
            'TARGET': 'Target',
            'AMAZON': 'Amazon',
            'COSTCO': 'Costco',
            'HOME DEPOT': 'Home Depot',
            'LOWES': 'Lowe\'s',
            'BEST BUY': 'Best Buy',
            'MACYS': 'Macy\'s',
            
            # Gas Stations
            'SHELL': 'Shell',
            'CHEVRON': 'Chevron',
            'EXXON': 'ExxonMobil',
            'BP': 'BP',
            'TEXACO': 'Texaco',
            'MOBIL': 'Mobil',
            
            # Banks & Financial
            'BANK OF AMERICA': 'Bank of America',
            'CHASE': 'JPMorgan Chase',
            'WELLS FARGO': 'Wells Fargo',
            'CITI': 'Citibank',
            'ATM FEE': 'ATM Fee',
            'OVERDRAFT': 'Overdraft Fee',
            
            # Utilities
            'ELECTRIC': 'Electric Company',
            'GAS COMPANY': 'Gas Company',
            'WATER DEPT': 'Water Department',
            'INTERNET': 'Internet Provider',
            'PHONE': 'Phone Company',
        }
        
        # Category mappings
        self.category_rules = {
            'Food & Dining': {
                'keywords': ['restaurant', 'food', 'dining', 'cafe', 'coffee', 'pizza', 'burger', 'fast food', 'delivery'],
                'vendors': ['McDonald\'s', 'KFC', 'Starbucks', 'Subway', 'Pizza Hut', 'Domino\'s Pizza', 'Burger King', 'Taco Bell', 'Chipotle'],
                'subcategories': {
                    'Fast Food': ['mcdonalds', 'kfc', 'burger king', 'taco bell', 'subway', 'fast food'],
                    'Coffee Shops': ['starbucks', 'coffee', 'cafe', 'espresso'],
                    'Restaurants': ['restaurant', 'dining', 'bistro', 'grill'],
                    'Delivery': ['delivery', 'doordash', 'ubereats', 'grubhub'],
                }
            },
            'Transportation': {
                'keywords': ['gas', 'fuel', 'parking', 'toll', 'uber', 'lyft', 'taxi', 'bus', 'train', 'metro'],
                'vendors': ['Shell', 'Chevron', 'ExxonMobil', 'BP', 'Texaco', 'Mobil'],
                'subcategories': {
                    'Gas & Fuel': ['gas', 'fuel', 'shell', 'chevron', 'exxon', 'bp', 'texaco', 'mobil'],
                    'Public Transit': ['bus', 'train', 'metro', 'subway', 'transit'],
                    'Ride Sharing': ['uber', 'lyft', 'taxi', 'cab'],
                    'Parking & Tolls': ['parking', 'toll', 'garage', 'meter'],
                }
            },
            'Shopping': {
                'keywords': ['store', 'shop', 'retail', 'purchase', 'buy', 'market', 'mall'],
                'vendors': ['Walmart', 'Target', 'Amazon', 'Costco', 'Home Depot', 'Lowe\'s', 'Best Buy', 'Macy\'s'],
                'subcategories': {
                    'General Merchandise': ['walmart', 'target', 'costco', 'store', 'retail'],
                    'Online Shopping': ['amazon', 'online', 'web', 'internet'],
                    'Home Improvement': ['home depot', 'lowes', 'hardware', 'garden'],
                    'Electronics': ['best buy', 'electronics', 'computer', 'phone'],
                    'Clothing': ['macys', 'clothing', 'apparel', 'fashion'],
                }
            },
            'Utilities': {
                'keywords': ['electric', 'electricity', 'gas', 'water', 'sewer', 'internet', 'phone', 'cable', 'utility'],
                'vendors': ['Electric Company', 'Gas Company', 'Water Department', 'Internet Provider', 'Phone Company'],
                'subcategories': {
                    'Electricity': ['electric', 'electricity', 'power', 'energy'],
                    'Gas': ['gas', 'natural gas', 'heating'],
                    'Water & Sewer': ['water', 'sewer', 'wastewater'],
                    'Internet & Phone': ['internet', 'phone', 'telecom', 'cable', 'wireless'],
                }
            },
            'Banking & Finance': {
                'keywords': ['fee', 'charge', 'interest', 'atm', 'overdraft', 'transfer', 'payment'],
                'vendors': ['Bank of America', 'JPMorgan Chase', 'Wells Fargo', 'Citibank', 'ATM Fee', 'Overdraft Fee'],
                'subcategories': {
                    'Bank Fees': ['fee', 'charge', 'atm', 'overdraft', 'maintenance'],
                    'Interest': ['interest', 'finance charge'],
                    'Transfers': ['transfer', 'wire', 'ach'],
                    'Payments': ['payment', 'autopay', 'bill pay'],
                }
            },
            'Healthcare': {
                'keywords': ['medical', 'doctor', 'hospital', 'pharmacy', 'health', 'dental', 'vision'],
                'vendors': [],
                'subcategories': {
                    'Medical': ['doctor', 'medical', 'hospital', 'clinic'],
                    'Pharmacy': ['pharmacy', 'prescription', 'drug', 'medicine'],
                    'Dental': ['dental', 'dentist', 'orthodontic'],
                    'Vision': ['vision', 'optical', 'eye', 'glasses'],
                }
            },
            'Entertainment': {
                'keywords': ['movie', 'theater', 'entertainment', 'streaming', 'netflix', 'spotify', 'game'],
                'vendors': [],
                'subcategories': {
                    'Streaming Services': ['netflix', 'hulu', 'spotify', 'streaming'],
                    'Movies & Theater': ['movie', 'theater', 'cinema'],
                    'Gaming': ['game', 'gaming', 'xbox', 'playstation'],
                    'Other Entertainment': ['entertainment', 'fun', 'leisure'],
                }
            },
            'Other': {
                'keywords': [],
                'vendors': [],
                'subcategories': {
                    'Miscellaneous': ['misc', 'other', 'unknown'],
                }
            }
        }
    
    def categorize_transaction(self, transaction: Dict) -> Tuple[str, str, str, str]:
        """
        Categorize a transaction and extract from/to information.
        
        Args:
            transaction: Transaction dictionary with raw_description
            
        Returns:
            Tuple of (from_account, to_account, expense_type, subcategory)
        """
        description = transaction.get('raw_description', '').lower().strip()
        
        # Extract vendor name using NER and mappings
        vendor_name = self._extract_vendor_name(description)
        
        # Categorize based on description and vendor
        expense_type, subcategory = self._categorize_by_rules(description, vendor_name)
        
        # Determine from/to accounts based on transaction type
        from_account, to_account = self._determine_accounts(transaction, vendor_name)
        
        return from_account, to_account, expense_type, subcategory
    
    def _extract_vendor_name(self, description: str) -> str:
        """Extract and normalize vendor name from description."""
        # First check direct mappings
        description_upper = description.upper()
        
        for key, vendor in self.vendor_mappings.items():
            if key in description_upper:
                return vendor
        
        # Use fuzzy matching for partial matches
        best_match = ""
        best_score = 0
        
        for key, vendor in self.vendor_mappings.items():
            score = fuzz.partial_ratio(key.lower(), description)
            if score > 80 and score > best_score:  # 80% similarity threshold
                best_match = vendor
                best_score = score
        
        if best_match:
            return best_match
        
        # Try NER if spaCy is available
        if self.nlp:
            try:
                doc = self.nlp(description)
                for ent in doc.ents:
                    if ent.label_ in ['ORG', 'PERSON']:
                        return ent.text.title()
            except:
                pass
        
        # Extract potential company name (usually first few words in caps)
        words = description.split()
        for i, word in enumerate(words[:3]):  # Check first 3 words
            if word.isupper() and len(word) > 2:
                potential_name = ' '.join(words[i:i+2]).title()
                return potential_name
        
        return "Unknown Vendor"
    
    def _categorize_by_rules(self, description: str, vendor_name: str) -> Tuple[str, str]:
        """Categorize transaction using rule-based approach."""
        best_category = "Other"
        best_subcategory = "Miscellaneous"
        best_score = 0
        
        for category, rules in self.category_rules.items():
            score = 0
            matched_subcategory = None
            
            # Check vendor matches
            if vendor_name in rules['vendors']:
                score += 50
            
            # Check keyword matches
            for keyword in rules['keywords']:
                if keyword in description:
                    score += 10
            
            # Check subcategory matches
            for subcategory, sub_keywords in rules['subcategories'].items():
                sub_score = 0
                for keyword in sub_keywords:
                    if keyword in description:
                        sub_score += 5
                if sub_score > 0 and (matched_subcategory is None or sub_score > best_score):
                    matched_subcategory = subcategory
                    score += sub_score
            
            if score > best_score:
                best_score = score
                best_category = category
                best_subcategory = matched_subcategory or list(rules['subcategories'].keys())[0]
        
        return best_category, best_subcategory
    
    def _determine_accounts(self, transaction: Dict, vendor_name: str) -> Tuple[str, str]:
        """Determine from and to accounts based on transaction type."""
        debit_amount = transaction.get('debit_amount', 0)
        credit_amount = transaction.get('credit_amount', 0)
        
        if debit_amount > 0:
            # Money going out - from user's account to vendor
            from_account = "My Account"
            to_account = vendor_name
        elif credit_amount > 0:
            # Money coming in - from vendor/source to user's account
            from_account = vendor_name
            to_account = "My Account"
        else:
            # Fallback
            from_account = "Unknown"
            to_account = "Unknown"
        
        return from_account, to_account

# main extract.py
"""
Main entry point for the bank statement processing pipeline.
"""
import argparse
import json
import logging
import sys
from pathlib import Path
from typing import List, Dict, Any

from file_loader import FileLoader
from preprocess import DataPreprocessor
from extractor import TransactionExtractor
from categorizer import TransactionCategorizer
from schema import Transaction, TransactionList

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('bank_statement_processor.log')
    ]
)
logger = logging.getLogger(__name__)

class BankStatementProcessor:
    """Main processor for bank statements."""
    
    def __init__(self):
        self.file_loader = FileLoader()
        self.preprocessor = DataPreprocessor()
        self.extractor = TransactionExtractor()
        self.categorizer = TransactionCategorizer()
    
    def process_file(self, file_path: str) -> TransactionList:
        """
        Process a bank statement file end-to-end.
        
        Args:
            file_path: Path to the bank statement file
            
        Returns:
            TransactionList object with validated transactions
        """
        logger.info(f"Starting processing of file: {file_path}")
        
        try:
            # Step 1: Load file
            file_type, raw_data = self.file_loader.load_file(file_path)
            
            # Step 2: Preprocess data
            if file_type in ['csv', 'excel']:
                processed_data = self.preprocessor.preprocess_structured_data(raw_data)
                raw_transactions = self.extractor.extract_from_structured_data(processed_data)
            else:  # pdf, docx
                processed_data = self.preprocessor.preprocess_text_data(raw_data)
                raw_transactions = self.extractor.extract_from_text_data(processed_data)
            
            # Step 3: Categorize transactions
            categorized_transactions = []
            for raw_transaction in raw_transactions:
                try:
                    from_acc, to_acc, expense_type, subcategory = self.categorizer.categorize_transaction(raw_transaction)
                    
                    # Create full transaction record
                    transaction_data = {
                        'from': from_acc,
                        'to': to_acc,
                        'credit_account': to_acc if raw_transaction['credit_amount'] > 0 else from_acc,
                        'debit_account': from_acc if raw_transaction['debit_amount'] > 0 else to_acc,
                        'expense_type': expense_type,
                        'subcategory': subcategory,
                        'debit_amount': raw_transaction['debit_amount'],
                        'credit_amount': raw_transaction['credit_amount'],
                        'transaction_date': raw_transaction['transaction_date']
                    }
                    
                    # Validate transaction
                    validated_transaction = Transaction(**transaction_data)
                    categorized_transactions.append(validated_transaction)
                    
                except Exception as e:
                    logger.warning(f"Error categorizing transaction: {str(e)}")
                    continue
            
            # Step 4: Create final result
            metadata = {
                'source_file': file_path,
                'file_type': file_type,
                'raw_transactions_found': len(raw_transactions),
                'valid_transactions': len(categorized_transactions),
                'processing_date': str(pd.Timestamp.now())
            }
            
            result = TransactionList(
                transactions=categorized_transactions,
                total_count=len(categorized_transactions),
                processing_metadata=metadata
            )
            
            logger.info(f"Successfully processed {len(categorized_transactions)} transactions")
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise

def main():
    """Main CLI entry point."""
    parser = argparse.ArgumentParser(description='Extract transactions from bank statements')
    parser.add_argument('file_path', help='Path to bank statement file')
    parser.add_argument('-o', '--output', help='Output JSON file path')
    parser.add_argument('-v', '--verbose', action='store_true', help='Enable verbose logging')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Validate input file
    if not Path(args.file_path).exists():
        print(f"Error: File not found - {args.file_path}")
        sys.exit(1)
    
    try:
        # Process the file
        processor = BankStatementProcessor()
        result = processor.process_file(args.file_path)
        
        # Prepare output
        output_data = result.dict()
        
        # Write to file or stdout
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                json.dump(output_data, f, indent=2, ensure_ascii=False)
            print(f"Results written to: {args.output}")
        else:
            print(json.dumps(output_data, indent=2, ensure_ascii=False))
        
        # Print summary
        print(f"\nSummary:")
        print(f"- Total transactions processed: {result.total_count}")
        print(f"- Source file: {result.processing_metadata['source_file']}")
        print(f"- File type: {result.processing_metadata['file_type']}")
        
        if result.total_count > 0:
            # Show category breakdown
            categories = {}
            for transaction in result.transactions:
                cat = transaction.expense_type
                categories[cat] = categories.get(cat, 0) + 1
            
            print(f"\nCategory Breakdown:")
            for category, count in sorted(categories.items()):
                print(f"- {category}: {count} transactions")
    
    except Exception as e:
        logger.error(f"Processing failed: {str(e)}")
        print(f"Error: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()

# ocr_processor.py (Optional OCR module for scanned PDFs)
"""
OCR processing for scanned bank statements using PaddleOCR.
"""
import logging
from typing import List, Tuple
import numpy as np
from pathlib import Path

try:
    from paddleocr import PaddleOCR
    import cv2
    import fitz  # PyMuPDF for PDF handling
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)

class OCRProcessor:
    """Handles OCR processing for scanned documents."""
    
    def __init__(self):
        if not OCR_AVAILABLE:
            raise ImportError("OCR dependencies not available. Install: pip install paddleocr PyMuPDF opencv-python")
        
        self.ocr = PaddleOCR(use_angle_cls=True, lang='en')
        self.logger = logging.getLogger(self.__class__.__name__)
    
    def process_scanned_pdf(self, pdf_path: str) -> List[str]:
        """
        Process scanned PDF using OCR.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            List of extracted text from each page
        """
        if not OCR_AVAILABLE:
            raise RuntimeError("OCR processing not available")
        
        self.logger.info(f"Processing scanned PDF with OCR: {pdf_path}")
        
        extracted_texts = []
        
        try:
            # Open PDF
            pdf_document = fitz.open(pdf_path)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                
                # Convert page to image
                mat = fitz.Matrix(2, 2)  # Scale factor for better OCR
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # Convert to numpy array for OpenCV
                nparr = np.frombuffer(img_data, np.uint8)
                img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                
                # Run OCR
                ocr_results = self.ocr.ocr(img, cls=True)
                
                # Extract text
                page_text = []
                if ocr_results and ocr_results[0]:
                    for result in ocr_results[0]:
                        text = result[1][0]  # Get text content
                        confidence = result[1][1]  # Get confidence score
                        
                        if confidence > 0.5:  # Filter low-confidence results
                            page_text.append(text)
                
                extracted_text = '\n'.join(page_text)
                extracted_texts.append(extracted_text)
                
                self.logger.info(f"Extracted {len(page_text)} text blocks from page {page_num + 1}")
            
            pdf_document.close()
            
        except Exception as e:
            self.logger.error(f"OCR processing failed: {str(e)}")
            raise
        
        return extracted_texts

# setup.py
"""
Setup script for the bank statement processor.
"""
from setuptools import setup, find_packages

setup(
    name="bank-statement-processor",
    version="1.0.0",
    description="Open-source bank statement processing pipeline",
    author="Your Name",
    author_email="your.email@example.com",
    packages=find_packages(),
    install_requires=[
        "pandas>=2.0.0",
        "pdfplumber>=0.9.0",
        "python-docx>=0.8.11",
        "paddlepaddle>=2.5.0",
        "paddleocr>=2.7.0",
        "spacy>=3.6.0",
        "pydantic>=2.0.0",
        "rapidfuzz>=3.0.0",
        "openpyxl>=3.1.0",
        "python-dateutil>=2.8.0",
        "PyMuPDF>=1.23.0",
        "opencv-python>=4.8.0",
    ],
    extras_require={
        "dev": [
            "pytest>=7.0.0",
            "black>=23.0.0",
            "isort>=5.12.0",
            "flake8>=6.0.0",
        ]
    },
    entry_points={
        "console_scripts": [
            "extract-transactions=extract:main",
        ]
    },
    python_requires=">=3.8",
    classifiers=[
        "Development Status :: 4 - Beta",
        "Intended Audience :: Financial and Insurance Industry",
        "License :: OSI Approved :: MIT License",
        "Programming Language :: Python :: 3.8",
        "Programming Language :: Python :: 3.9",
        "Programming Language :: Python :: 3.10",
        "Programming Language :: Python :: 3.11",
    ],
)

# README.md
"""
# Bank Statement Processing Pipeline

A comprehensive, open-source solution for extracting and categorizing transactions from bank statements in multiple formats (CSV, XLSX, PDF, DOCX).

## Features

- **Multi-format Support**: Handles CSV, Excel, PDF, and Word documents
- **OCR Processing**: Extracts text from scanned PDFs using PaddleOCR
- **Smart Categorization**: Uses NLP and rule-based approaches for transaction categorization
- **Vendor Recognition**: Maps common abbreviations to full vendor names
- **Data Validation**: Pydantic schemas ensure output consistency
- **Modular Architecture**: Clean, maintainable code structure

## Installation

1. **Clone or download the code**
```bash
# Save all Python files in a project directory
mkdir bank-statement-processor
cd bank-statement-processor
```

2. **Install Python dependencies**
```bash
pip install pandas pdfplumber python-docx paddlepaddle paddleocr spacy pydantic rapidfuzz openpyxl python-dateutil PyMuPDF opencv-python
```

3. **Download spaCy language model**
```bash
python -m spacy download en_core_web_sm
```

## Usage

### Command Line Interface

```bash
# Basic usage
python extract.py path/to/statement.pdf

# Save output to file
python extract.py path/to/statement.xlsx -o transactions.json

# Verbose logging
python extract.py path/to/statement.csv -v
```

### Python API

```python
from extract import BankStatementProcessor

processor = BankStatementProcessor()
result = processor.process_file('statement.pdf')

# Access transactions
for transaction in result.transactions:
    print(f"{transaction.transaction_date}: {transaction.expense_type} - ${transaction.debit_amount}")
```

## Output Format

```json
{
  "transactions": [
    {
      "from": "My Account",
      "to": "McDonald's",
      "credit_account": "My Account",
      "debit_account": "McDonald's",
      "expense_type": "Food & Dining",
      "subcategory": "Fast Food",
      "debit_amount": 12.50,
      "credit_amount": 0.0,
      "transaction_date": "2024-01-15"
    }
  ],
  "total_count": 1,
  "processing_metadata": {
    "source_file": "statement.pdf",
    "file_type": "pdf",
    "raw_transactions_found": 1,
    "valid_transactions": 1,
    "processing_date": "2024-01-20 10:30:00"
  }
}
```

## Supported File Formats

- **CSV**: Comma-separated values with automatic encoding detection
- **Excel**: XLSX and XLS files with multi-sheet support
- **PDF**: Both text-based and scanned PDFs (with OCR)
- **Word**: DOCX files with table and paragraph extraction

## Architecture

- `extract.py` - Main entry point and CLI
- `file_loader.py` - File format handling
- `preprocess.py` - Data cleaning and normalization
- `extractor.py` - Transaction extraction logic
- `categorizer.py` - NLP-based categorization
- `schema.py` - Data validation schemas
- `ocr_processor.py` - OCR processing for scanned documents

## Categories

The system automatically categorizes transactions into:

- Food & Dining (Fast Food, Coffee Shops, Restaurants, Delivery)
- Transportation (Gas & Fuel, Public Transit, Ride Sharing, Parking & Tolls)
- Shopping (General Merchandise, Online Shopping, Home Improvement, Electronics, Clothing)
- Utilities (Electricity, Gas, Water & Sewer, Internet & Phone)
- Banking & Finance (Bank Fees, Interest, Transfers, Payments)
- Healthcare (Medical, Pharmacy, Dental, Vision)
- Entertainment (Streaming Services, Movies & Theater, Gaming)
- Other (Miscellaneous)

## Error Handling

The system includes comprehensive error handling for:
- File format issues
- Encoding problems
- Date parsing errors
- Invalid amounts
- Missing data

## Logging

All processing steps are logged with timestamps for debugging and monitoring.

## License

This project is open-source and available under the MIT License.
"""