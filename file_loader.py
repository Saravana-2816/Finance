import os
import logging
from typing import Union, Tuple, List
import pandas as pd
import pdfplumber
from docx import Document
from pathlib import Path

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileLoader:
    """Handles loading of various file formats."""
    
    SUPPORTED_EXTENSIONS = {'.csv', '.xlsx', '.xls', '.pdf', '.docx'}
    
    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)
    
    def load_file(self, file_path: str) -> Tuple[str, Union[pd.DataFrame, str, List[str]]]:
        """
        Load file based on extension and return (file_type, content).
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (file_type, content) where content varies by type
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        self.logger.info(f"Loading {file_ext} file: {file_path}")
        
        try:
            if file_ext == '.csv':
                return self._load_csv(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._load_excel(file_path)
            elif file_ext == '.pdf':
                return self._load_pdf(file_path)
            elif file_ext == '.docx':
                return self._load_docx(file_path)
        except Exception as e:
            self.logger.error(f"Error loading file {file_path}: {str(e)}")
            raise
    
    def _load_csv(self, file_path: str) -> Tuple[str, pd.DataFrame]:
        """Load CSV file with robust encoding detection."""
        encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']
        
        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, encoding=encoding, low_memory=False)
                self.logger.info(f"Successfully loaded CSV with {encoding} encoding")
                return 'csv', df
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode CSV file with any of the tried encodings: {encodings}")
    
    def _load_excel(self, file_path: str) -> Tuple[str, pd.DataFrame]:
        """Load Excel file, trying all sheets."""
        try:
            # Try to read all sheets and combine
            excel_file = pd.ExcelFile(file_path)
            
            # If only one sheet, return it directly
            if len(excel_file.sheet_names) == 1:
                df = pd.read_excel(file_path, sheet_name=0)
                return 'excel', df
            
            # Multiple sheets - look for the one most likely to contain transactions
            transaction_keywords = ['transaction', 'statement', 'activity', 'history']
            
            for sheet_name in excel_file.sheet_names:
                if any(keyword in sheet_name.lower() for keyword in transaction_keywords):
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    self.logger.info(f"Using sheet: {sheet_name}")
                    return 'excel', df
            
            # If no obvious sheet found, use the first one
            df = pd.read_excel(file_path, sheet_name=0)
            self.logger.info(f"Using first sheet: {excel_file.sheet_names[0]}")
            return 'excel', df
            
        except Exception as e:
            raise ValueError(f"Error reading Excel file: {str(e)}")
    
    def _load_pdf(self, file_path: str) -> Tuple[str, List[str]]:
        """Extract text from PDF using pdfplumber."""
        pages_text = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        pages_text.append(text)
                        self.logger.info(f"Extracted text from page {i+1}")
            
            if not pages_text:
                self.logger.warning("No text extracted from PDF - may need OCR")
            
            return 'pdf', pages_text
            
        except Exception as e:
            raise ValueError(f"Error reading PDF file: {str(e)}")
    
    def _load_docx(self, file_path: str) -> Tuple[str, str]:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))
            
            combined_text = '\n'.join(full_text)
            self.logger.info(f"Extracted {len(full_text)} text blocks from DOCX")
            
            return 'docx', combined_text
            
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {str(e)}")